"""Convert CYB302_Lab_Report.md to a formatted .docx file."""

import re
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

BASE_DIR = Path(__file__).resolve().parent.parent
MD_PATH = BASE_DIR / "CYB302_Lab_Report.md"
DOCX_PATH = BASE_DIR / "CYB302_Lab_Report.docx"


def add_code_block(doc, code):
    """Add a code block styled with monospace font and grey background."""
    p = doc.add_paragraph()
    p.style = doc.styles['No Spacing']
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(code)
    run.font.name = 'Consolas'
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor(0x1E, 0x1E, 0x1E)
    # Add shading
    shading = run._element.get_or_add_rPr()
    shd = shading.makeelement(qn('w:shd'), {
        qn('w:fill'): 'F2F2F2',
        qn('w:val'): 'clear'
    })
    shading.append(shd)


def add_table(doc, md_table):
    """Parse a markdown table and add it to the document."""
    lines = [l.strip() for l in md_table.strip().split('\n') if l.strip()]
    if len(lines) < 2:
        return
    header_line = lines[0]
    rows_data = lines[2:]

    # Parse header
    headers = [h.strip() for h in header_line.split('|')[1:-1]]

    rows = []
    for row_line in rows_data:
        cols = [c.strip() for c in row_line.split('|')[1:-1]]
        if cols:
            rows.append(cols)

    if not headers:
        return

    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(9)

    # Data rows
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            if c_idx < len(headers):
                cell = table.rows[r_idx + 1].cells[c_idx]
                cell.text = val
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(9)
    doc.add_paragraph()


def add_image_placeholder(doc, text):
    """Add a screenshot placeholder."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f'[ {text} ]')
    run.bold = True
    run.font.color.rgb = RGBColor(0x00, 0x70, 0xC0)
    run.font.size = Pt(10)
    run.italic = True


def parse_and_build():
    with open(MD_PATH, 'r', encoding='utf-8') as f:
        content = f.read()

    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.9)
        section.right_margin = Inches(0.9)

    # Title page
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(120)
    run = p.add_run('CYB 302 — Biometrics Security')
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(0x1A, 0x47, 0x8A)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Lab Report: Face Recognition Authentication System')
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(30)
    run = p.add_run('Course: CYB 302 — Biometrics Security\n'
                    'Project: Face Recognition for University Secure Lab Access\n'
                    'Technology: FastAPI + Python + OpenCV + dlib + SQLite\n'
                    'Dataset: ORL Face Database — 40 subjects, 10 images each')
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    doc.add_page_break()

    # Process content
    lines = content.split('\n')
    i = 0
    in_code_block = False
    in_table = False
    table_lines = []
    skip_until_section = False

    while i < len(lines):
        line = lines[i]

        # Code blocks (```)
        if line.strip().startswith('```'):
            if in_code_block:
                in_code_block = False
                i += 1
                continue
            else:
                in_code_block = True
                code_content = []
                i += 1
                while i < len(lines) and not lines[i].strip().startswith('```'):
                    code_content.append(lines[i])
                    i += 1
                add_code_block(doc, '\n'.join(code_content))
                in_code_block = False
                i += 1
                continue

        # Skip table of contents and screenshot placeholders inside code blocks
        # Tables
        if line.startswith('|') and line.endswith('|') and not line.startswith('|---'):
            table_lines.append(line)
            in_table = True
            i += 1
            continue

        if in_table:
            if line.startswith('|---'):
                i += 1
                continue
            if line.startswith('|') and line.endswith('|'):
                table_lines.append(line)
                i += 1
                continue
            else:
                add_table(doc, '\n'.join(table_lines))
                table_lines = []
                in_table = False
                continue

        stripped = line.strip()

        # Skip empty lines after certain elements
        if not stripped:
            i += 1
            continue

        # Screenshot placeholders
        if stripped.startswith('**[SCREENSHOT:') and stripped.endswith(']**'):
            text = stripped.replace('**[SCREENSHOT:', '').replace(']**', '').strip()
            add_image_placeholder(doc, f'Screenshot: {text}')
            doc.add_paragraph()
            i += 1
            continue

        # Horizontal rule
        if stripped == '---':
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            run = p.add_run('_' * 70)
            run.font.color.rgb = RGBColor(0xBB, 0xBB, 0xBB)
            run.font.size = Pt(8)
            i += 1
            continue

        # Headings
        if stripped.startswith('## '):
            text = stripped[3:].strip()
            # Remove bold markers
            text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(16)
            p.paragraph_format.space_after = Pt(6)
            run = p.add_run(text)
            run.bold = True
            run.font.size = Pt(14)
            run.font.color.rgb = RGBColor(0x1A, 0x47, 0x8A)
            i += 1
            continue

        if stripped.startswith('### '):
            text = stripped[4:].strip()
            text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(4)
            run = p.add_run(text)
            run.bold = True
            run.font.size = Pt(12)
            run.font.color.rgb = RGBColor(0x2C, 0x5F, 0x9A)
            i += 1
            continue

        if stripped.startswith('#### '):
            text = stripped[5:].strip()
            text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(3)
            run = p.add_run(text)
            run.bold = True
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor(0x3A, 0x6F, 0xAA)
            i += 1
            continue

        # Regular paragraph with inline formatting
        # Process bold **text**
        parts = re.split(r'(\*\*.*?\*\*)', stripped)
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)

        # Check if it's an indented code line (starts with 4 spaces for code in list)
        if stripped.startswith('    '):
            run = p.add_run(stripped)
            run.font.name = 'Consolas'
            run.font.size = Pt(8.5)
            run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
            i += 1
            continue

        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                run = p.add_run(part[2:-2])
                run.bold = True
                run.font.size = Pt(10.5)
            else:
                run = p.add_run(part)
                run.font.size = Pt(10.5)

        i += 1

    # Save
    doc.save(str(DOCX_PATH))
    print(f"Report saved to: {DOCX_PATH}")


if __name__ == '__main__':
    parse_and_build()
